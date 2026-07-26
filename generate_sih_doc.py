import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_verbatim_pdf_template_doc(output_path="Honeywell_Idea_Submission_BehaviorIQ.docx"):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    def add_page_title(text):
        p = doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42) # Dark Slate
        return p

    def add_main_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(11.5)
        return p

    def add_detail_bullet(bold_prefix, text=""):
        p = doc.add_paragraph(style='List Bullet 2')
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.font.bold = True
            r1.font.size = Pt(10.5)
        if text:
            r2 = p.add_run(" " + text if bold_prefix else text)
            r2.font.size = Pt(10.5)
        return p

    def add_footer_tag(page_num):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"@SIH Idea submission- Template {page_num}")
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(148, 163, 184) # Light Slate

    # ================= PAGE 1 =================
    add_page_title("IMPORTANT INSTRUCTIONS")
    doc.add_paragraph()
    p1 = doc.add_paragraph("Please ensure below pointers are met while submitting the Idea PPT:")
    p1.runs[0].font.bold = True
    
    doc.add_paragraph("1. Kindly keep the maximum slides limit up to six (6). ( Including the title slide)", style='List Bullet')
    doc.add_paragraph("2. Try to avoid paragraphs and post your idea in points /diagrams / Infographics /pictures", style='List Bullet')
    doc.add_paragraph("3. Keep your explanation precise and easy to understand", style='List Bullet')
    doc.add_paragraph("4. Idea should be unique and novel.", style='List Bullet')
    doc.add_paragraph("5. You can only use provided template for making the PPT without changing the idea details pointers (mentioned in previous slides).", style='List Bullet')
    doc.add_paragraph("6. You need to save the file in PDF and upload the same on portal. No PPT, Word Doc or any other format will be supported.", style='List Bullet')
    
    doc.add_paragraph()
    p_note = doc.add_paragraph()
    r_note = p_note.add_run("Note - You can delete this slide (Important Pointers) when you upload the details of your idea portal.")
    r_note.font.bold = True
    r_note.font.color.rgb = RGBColor(220, 38, 38)
    
    add_footer_tag(1)
    doc.add_page_break()

    # ================= PAGE 2 =================
    add_page_title("TITLE PAGE")
    doc.add_paragraph()
    add_main_bullet("Problem Statement ID – [Insert Your PS ID e.g. SIH1234 / HW_SEC_01]")
    add_main_bullet("Problem Statement Title- Autonomous Behavioral Anomaly Detection & Threat SOC Engine for Enterprise IT & Industrial OT")
    add_main_bullet("Theme- Cybersecurity / AI & ML / Smart Industrial Automation")
    add_main_bullet("PS Category- Software")
    add_main_bullet("Student Name (Registered on portal)- Yash Singh")
    add_main_bullet("Student ID- [Insert Your Registered Portal Student ID]")
    
    add_footer_tag(2)
    doc.add_page_break()

    # ================= PAGE 3 =================
    add_page_title("IDEA TITLE: BehaviorIQ")
    doc.add_paragraph()
    
    p_heading = doc.add_paragraph()
    r_h = p_heading.add_run("Proposed Solution (Describe your Idea/Solution/Prototype)")
    r_h.font.bold = True
    r_h.font.size = Pt(13)
    r_h.font.underline = True

    add_main_bullet("Detailed explanation of the proposed solution")
    add_detail_bullet("• BehaviorIQ Engine:", "Autonomous AI-powered Security Operations Center (SOC) engine replacing static SIEM alert rules with dynamic, per-entity machine learning baselines.")
    add_detail_bullet("• Dual-Stage ML Core:", "Unsupervised Isolation Forest (5,000 decision trees) for anomaly scoring [0-100] paired with a Multi-Class LightGBM Classifier for threat taxonomy assignment.")
    add_detail_bullet("• Explainable AI (SHAP XAI):", "Calculates SHAP feature attributions and generates natural-language forensic narratives for every flagged alert.")

    add_main_bullet("How it addresses the problem")
    add_detail_bullet("• 99% Noise Reduction:", "Enforces a strict Top 1% Alert Budget (>= 75.0 Risk Score), eliminating analyst alert fatigue.")
    add_detail_bullet("• Sub-Second Mitigation:", "Executes autonomous playbooks (Cloudflare WAF IP quarantine, OAuth2 session token revocation, mandatory MFA) in <0.38 seconds.")
    add_detail_bullet("• Edge Case Mastery:", "Solves Cold-Start entities (N < 20 events) via Hierarchical Bayesian Blending (0.85% FPR) and Concept Drift via EMA (alpha = 0.05).")

    add_main_bullet("Innovation and uniqueness of the solution")
    add_detail_bullet("• Real-Time Geo-Velocity:", "Calculates physical travel speed (> 840 km/h) between consecutive access tokens to stop impossible travel attacks.")
    add_detail_bullet("• Dark Web Intelligence:", "Live cross-referencing of entity tokens against active dark web credential dumps.")

    add_footer_tag(3)
    doc.add_page_break()

    # ================= PAGE 4 =================
    add_page_title("TECHNICAL APPROACH")
    doc.add_paragraph()

    add_main_bullet("Technologies to be used (e.g. programming languages, frameworks, hardware)")
    add_detail_bullet("• Programming & Frameworks:", "Python 3.11, FastAPI REST Framework, Uvicorn, Gunicorn WSGI.")
    add_detail_bullet("• ML & Explainability:", "Scikit-Learn (Isolation Forest), LightGBM (Multi-Class Gradient Boosting), SHAP (XAI), NumPy, Pandas.")
    add_detail_bullet("• Ingestion & Cache:", "Apache Kafka Event Broker (1,420+ eps), Redis Feature Store Cache (<0.2ms latency).")
    add_detail_bullet("• Frontend Interface:", "Glassmorphism HTML5 Canvas, Tailwind CSS v3, Chart.js, Web Audio API.")
    add_detail_bullet("• Microservices & Cloud:", "Docker, Docker Compose, Kubernetes HPA, Vercel (Frontend), Render (Backend).")

    add_main_bullet("Methodology and process for implementation (Flow Charts/Images/ working prototype)")
    add_detail_bullet("1. Ingestion & Feature Extraction:", "Access Logs Stream (1,420 eps) ➔ Extract Geo-Velocity, Failed Login Bursts, Device Fingerprints.")
    add_detail_bullet("2. Profiling & Scoring:", "Feature Matrix ➔ Isolation Forest Profiler (5,000 Trees) ➔ Anomaly Score [0-100].")
    add_detail_bullet("3. Alert Budget Filter:", "Risk Score >= 75.0 ? Flag Threat : Update EMA Baseline Profile (alpha = 0.05).")
    add_detail_bullet("4. Classification & XAI:", "Flagged Incident ➔ LightGBM Classifier ➔ SHAP Feature Importance Attribution ➔ Triage Desk.")
    add_detail_bullet("5. Autonomous SOAR:", "High Confidence Threat ➔ Execute Cloudflare WAF Block & OAuth Token Revocation (<0.38s).")

    add_footer_tag(4)
    doc.add_page_break()

    # ================= PAGE 5 =================
    add_page_title("FEASIBILITY AND VIABILITY")
    doc.add_paragraph()

    add_main_bullet("Analysis of the feasibility of the idea")
    add_detail_bullet("• ML Benchmark Accuracy:", "Binary Anomaly F1-Score = 0.961, Precision = 0.972, Recall = 0.951.")
    add_detail_bullet("• Real-Time Latency:", "Inference completes in <0.8ms per event, handling >12,000 events/sec throughput.")
    add_detail_bullet("• Alert Budget Control:", "Top 1% cutoff maintains False Positive Rate < 0.32%.")

    add_main_bullet("Potential challenges and risks")
    add_detail_bullet("• Cold-Start Challenge:", "Newly onboarded users or IoT edge devices (N < 20 events) lacking historical baselines, risking false positive storms.")
    add_detail_bullet("• Concept Drift Risk:", "Evolving employee shift patterns causing baseline drift over time.")

    add_main_bullet("Strategies for overcoming these challenges")
    add_detail_bullet("• Hierarchical Bayesian Blending:", "Interpolates population priors (user vs service_account vs edge_device) with empirical stats, reducing cold-start FPR to 0.85%.")
    add_detail_bullet("• EMA Baseline Adaptation (alpha = 0.05):", "Dynamically updates non-anomalous entity centroids over time without requiring expensive model retraining.")

    add_footer_tag(5)
    doc.add_page_break()

    # ================= PAGE 6 =================
    add_page_title("ARTIFACTS")
    doc.add_paragraph()

    add_main_bullet("Relevant artifacts, such as :")

    add_main_bullet("Copy of the Code Embedded")
    add_detail_bullet("• GitHub Repository:", "https://github.com/Yash-Singh607/Behavior-IQ")
    add_detail_bullet("• Production Containers:", "Dockerfile (512MB RAM optimized), docker-compose.yml (API, Redis, Kafka).")
    add_detail_bullet("• Kubernetes Manifests:", "deploy/k8s/deployment.yaml (Deployment, Service LoadBalancer, Horizontal Pod Autoscaler).")

    add_main_bullet("Snaps of the solution proposal")
    add_detail_bullet("• Live Vercel Frontend UI:", "https://behavior-iq-jade.vercel.app")
    add_detail_bullet("• Live Render FastAPI Backend:", "https://behavior-iq.onrender.com")
    add_detail_bullet("• Interactive Swagger API Docs:", "https://behavior-iq.onrender.com/docs")
    add_detail_bullet("• Analyst Login Credentials:", "Email: admin@behavioriq.ai | Password: admin123")

    add_main_bullet("Dashboard snaps")
    add_detail_bullet("• Incident Triage Workbench:", "Real-time triage stream table with SHAP feature attributions.")
    add_detail_bullet("• Stage Threat Injector:", "⚡ Launch Attack Scenario trigger (APT29 Cozy Bear, FIN7 Ransomware, Insider Data Theft).")
    add_detail_bullet("• Security Gate Authentication:", "Glassmorphism login gate modal with user profile dropdown.")

    add_footer_tag(6)
    doc.add_page_break()

    # ================= PAGE 7 =================
    add_page_title("RESEARCH AND REFERENCES")
    doc.add_paragraph()

    add_main_bullet("Details / Links of the reference and research work")
    add_detail_bullet("• Isolation Forest Anomaly Detection:", "Liu, F. T., Ting, K. M., & Zhou, Z. H. (2008). Isolation Forest. IEEE International Conference on Data Mining (ICDM).")
    add_detail_bullet("• LightGBM Gradient Boosting Framework:", "Ke, G., et al. (2017). LightGBM: A Highly Efficient Gradient Boosting Decision Tree. Advances in Neural Information Processing Systems (NeurIPS).")
    add_detail_bullet("• SHAP Explainable AI Framework:", "Lundberg, S. M., & Lee, S. I. (2017). A Unified Approach to Interpreting Model Predictions. NeurIPS.")
    add_detail_bullet("• MITRE ATT&CK Framework:", "MITRE ATT&CK Matrix for Enterprise & Industrial Control Systems (ICS).")
    add_detail_bullet("• Compliance Standards:", "ISO/IEC 27001:2022, SOC 2 Type II, and GDPR Privacy Controls.")

    add_footer_tag(7)

    doc.save(output_path)
    print(f"Verbatim template document successfully created at {output_path}")

if __name__ == "__main__":
    create_verbatim_pdf_template_doc()

